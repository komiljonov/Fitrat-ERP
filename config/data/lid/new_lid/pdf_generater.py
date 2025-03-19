import io
import os
from django.http import HttpResponse
from rest_framework.views import APIView
from django.utils.timezone import now
from drf_yasg.utils import swagger_auto_schema
from drf_yasg import openapi
from docx import Document
from docx.shared import Inches


SIGNATURE_PATH = os.path.abspath("data/lid/new_lid/sign.png")
STAMP_PATH = os.path.abspath("data/lid/new_lid/pechate.png")
TEMPLATE_PATH = os.path.abspath("data/lid/new_lid/Shartnoma.docx")

def replace_image(doc, placeholder, image_path, width=1.3):
    """Replace a placeholder with an image inside tables with correct size."""
    if not os.path.exists(image_path):
        print(f"❌ Image Not Found: {image_path}")  # Debugging output
        return

    found_placeholder = False

    # ✅ Replace in table cells (Ensures correct positioning)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if placeholder in para.text:
                        found_placeholder = True
                        para.text = para.text.replace(placeholder, "")
                        run = para.add_run()
                        run.add_picture(image_path, width=Inches(width))
                        para.alignment = 1  # Center align

    if not found_placeholder:
        print(f"⚠️ Warning: Placeholder '{placeholder}' not found in document.")

class ContractGenerateAPIView(APIView):
    @swagger_auto_schema(
        request_body=openapi.Schema(
            type=openapi.TYPE_OBJECT,
            required=["student__full_name", "subjects", "from_date", "from_month", "from_year", "to_date", "to_month", "to_year"],
            properties={
                "student__full_name": openapi.Schema(type=openapi.TYPE_STRING, description="Full name of the student"),
                "subjects": openapi.Schema(type=openapi.TYPE_STRING, description="Subjects the student is enrolled in"),
                "from_date": openapi.Schema(type=openapi.TYPE_INTEGER, description="Start date (day)"),
                "from_month": openapi.Schema(type=openapi.TYPE_STRING, description="Start month"),
                "from_year": openapi.Schema(type=openapi.TYPE_INTEGER, description="Start year"),
                "to_date": openapi.Schema(type=openapi.TYPE_INTEGER, description="End date (day)"),
                "to_month": openapi.Schema(type=openapi.TYPE_STRING, description="End month"),
                "to_year": openapi.Schema(type=openapi.TYPE_INTEGER, description="End year"),
            },
        ),
        responses={200: "Returns a contract file for download"},
    )
    def post(self, request):
        """Generate a contract document with properly positioned sign & stamp images."""
        data = request.data
        student_name = data.get("student__full_name", "____")
        subjects = data.get("subjects", "____")
        from_date = data.get("from_date", now().day)
        from_month = data.get("from_month", now().strftime("%B"))
        from_year = data.get("from_year", now().year)
        to_date = data.get("to_date", now().day)
        to_month = data.get("to_month", now().strftime("%B"))
        to_year = data.get("to_year", now().year + 1)
        day = data.get("day", now().day)
        month = data.get("month", now().strftime("%B"))

        # ✅ Load the contract template
        doc = Document(TEMPLATE_PATH)

        # ✅ Replace text placeholders
        for para in doc.paragraphs:
            para.text = para.text.replace("{student__full_name}", student_name)
            para.text = para.text.replace("{subjects}", subjects)
            para.text = para.text.replace("{from_date}", str(from_date))
            para.text = para.text.replace("{from_month}", from_month)
            para.text = para.text.replace("{from_year}", str(from_year))
            para.text = para.text.replace("{to_date}", str(to_date))
            para.text = para.text.replace("{to_month}", to_month)
            para.text = para.text.replace("{to_year}", str(to_year))
            para.text = para.text.replace("{day}", str(day))
            para.text = para.text.replace("{month}", month)

        # ✅ Reduce `{sign}` and `{pechate}` sizes to fit in their placeholders
        replace_image(doc, "{sign}", SIGNATURE_PATH, width=1.3)  # Smaller signature
        replace_image(doc, "{pechate}", STAMP_PATH, width=1.5)  # Smaller stamp

        # ✅ Save the document to memory
        output_stream = io.BytesIO()
        doc.save(output_stream)
        output_stream.seek(0)

        # ✅ Return the file as a response
        response = HttpResponse(output_stream, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        response["Content-Disposition"] = f"attachment; filename=contract_{student_name}.docx"
        return response
